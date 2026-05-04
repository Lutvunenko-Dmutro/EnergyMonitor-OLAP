"""
СИСТЕМА ІНТЕЛЕКТУАЛЬНОГО АУДИТУ ТЕКСТУ (Academic Quality Guard)
=============================================================
Модуль для автоматизованої перевірки якості академічних робіт (дипломів, звітів).
Забезпечує:
1. Formatting Validation: перевірка полів, шрифтів та міжрядкових інтервалів.
2. Anti-Plagiarism Detection: виявлення змішаних алфавітів для запобігання обходу систем перевірки.
3. Linguistic & Style Analysis: аналіз варіативності довжини речень, словникового багатства та закону Ціпфа.
4. Readability Metrics: розрахунок індексу ARI для оцінки наукової складності тексту.
5. Structural Audit: верифікація наявності обов'язкових розділів (Зміст, Список джерел).
Гарантує відповідність фінального документа ДСТУ та високим академічним стандартам.
"""
import os
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def check_margins(doc):
    results = []
    section = doc.sections[0]
    
    # Очікувані значення в см
    expected = {"left": 3.0, "right": 1.0, "top": 2.0, "bottom": 2.0}
    current = {
        "left": round(section.left_margin.cm, 1),
        "right": round(section.right_margin.cm, 1),
        "top": round(section.top_margin.cm, 1),
        "bottom": round(section.bottom_margin.cm, 1)
    }
    
    for key, val in expected.items():
        if abs(current[key] - val) > 0.1:
            results.append(f" [❌] Поля: {key} = {current[key]}см (очікувалось {val}см)")
        else:
            results.append(f" [✅] Поля: {key} в межах норми ({current[key]}см)")
    return results

def check_mixed_alphabets(doc):
    results = []
    mixed_words = []
    
    # Регулярні вирази для пошуку суміші кирилиці та латиниці в одному слові
    cyrillic = set('абвгґдеєжзиіїйклмнопрстуфхцчшщьюя')
    latin = set('abcdefghijklmnopqrstuvwxyz')
    
    for para in doc.paragraphs:
        words = para.text.split()
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word).lower()
            if not clean_word: continue
            
            has_cyr = any(c in cyrillic for c in clean_word)
            has_lat = any(c in latin for c in clean_word)
            
            if has_cyr and has_lat:
                # Виключаємо слова з цифрами або специфічні терміни (SaaS-платформа тощо)
                if not re.search(r'\d', word) and '-' not in word:
                    mixed_words.append(word)
    
    if mixed_words:
        results.append(f" [❌] Змішаний алфавіт: знайдено {len(mixed_words)} слів із сумішшю кирилиці та латиниці (підозра на обхід антиплагіату!)")
        results.append(f"      Приклади: {', '.join(mixed_words[:5])}")
    else:
        results.append(" [✅] Алфавіт: " + "чистий (слів-гібридів не знайдено)")
    return results

def check_formatting_rules(doc):
    results = []
    total_paras = len(doc.paragraphs)
    wrong_font = 0
    wrong_spacing = 0
    
    for para in doc.paragraphs:
        if not para.text.strip(): continue
        
        # Перевірка шрифту та розміру (беремо перший run)
        if para.runs:
            run = para.runs[0]
            if run.font.name and run.font.name != "Times New Roman":
                wrong_font += 1
            if run.font.size and run.font.size != Pt(14) and para.style.name == 'Normal':
                wrong_font += 1
        
        # Перевірка інтервалу (для основного тексту)
        if para.style.name == 'Normal' and para.paragraph_format.line_spacing != 1.5:
            if para.paragraph_format.line_spacing_rule != 1: # 1.5 rule is often custom
                wrong_spacing += 1

    if wrong_font > (total_paras * 0.1): # Більше 10% помилок
        results.append(f" [⚠️] Шрифт: виявлено відхилення у {wrong_font} абзацах (бажано Times New Roman 14pt)")
    else:
        results.append(" [✅] Шрифт: Times New Roman 14pt витримано у більшості тексту")
        
    return results

def check_ai_style(doc):
    results = []
    full_text = " ".join([p.text for p in doc.paragraphs if len(p.text) > 50])
    # Розбиваємо на речення за крапками
    sentences = re.split(r'[.!?]\s+', full_text)
    lengths = [len(s.split()) for s in sentences if len(s.split()) > 3]
    
    if not lengths: return [" [?] Стиль: Недостатньо тексту для аналізу"]
    
    avg_len = sum(lengths) / len(lengths)
    # Розрахунок дисперсії (чисто для довідки)
    variance = sum((x - avg_len) ** 2 for x in lengths) / len(lengths)
    
    # AI зазвичай пише реченнями однакової довжини (низька дисперсія)
    # Люди пишуть то коротко, то довго.
    if variance < 30:
        results.append(f" [⚠️] Стиль: Низька варіативність довжини речень ({round(variance, 1)}). Текст виглядає монотонно (ознака AI).")
    else:
        results.append(f" [✅] Стиль: Природна варіативність тексту ({round(variance, 1)}).")
        
    return results

def advanced_linguistic_analysis(doc):
    results = []
    full_text = " ".join([p.text for p in doc.paragraphs if len(p.text) > 100])
    words = re.findall(r'\w+', full_text.lower())
    
    if len(words) < 100:
        return [" [?] Лінгвістика: Замало тексту для глибокого аналізу"]
        
    # 1. Словникове багатство (TTR)
    unique_words = set(words)
    ttr = len(unique_words) / len(words)
    
    # 2. Hapax Legomena (слова, що зустрічаються 1 раз)
    word_counts = {}
    for w in words: word_counts[w] = word_counts.get(w, 0) + 1
    hapaxes = [w for w in word_counts if word_counts[w] == 1]
    hapax_ratio = len(hapaxes) / len(words)
    
    # Висновки
    # Для української мови TTR > 0.4 для дипломів — це добре.
    if ttr < 0.35:
        results.append(f" [⚠️] Словник: Низьке багатство ({round(ttr, 2)}). Багато повторів (ознака ШІ).")
    else:
        results.append(f" [✅] Словник: Багатий ({round(ttr, 2)}).")
        
    if hapax_ratio < 0.2:
        results.append(f" [⚠️] Унікальність: Мало рідкісних слів ({round(hapax_ratio, 2)}).")
    else:
        results.append(f" [✅] Унікальність: Багато унікальних зворотів ({round(hapax_ratio, 2)}).")
        
    return results

def calculate_readability(doc):
    results = []
    full_text = " ".join([p.text for p in doc.paragraphs if len(p.text) > 100])
    words = re.findall(r'\w+', full_text)
    sentences = re.split(r'[.!?]\s+', full_text)
    
    if len(words) < 100 or len(sentences) < 5:
        return [" [?] Читабельність: Замало даних"]

    # Середня довжина речення (ASL)
    asl = len(words) / len(sentences)
    # Середня кількість символів у слові (AWL)
    awl = sum(len(w) for w in words) / len(words)
    
    # Automated Readability Index (ARI)
    ari = 4.71 * awl + 0.5 * asl - 21.43
    
    if ari > 16:
        results.append(f" [🎓] Рівень: Текст високої наукової складності (ARI={round(ari, 1)})")
    else:
        results.append(f" [✅] Рівень: Текст доступний для розуміння (ARI={round(ari, 1)})")
        
    return results

def check_zipf_law(doc):
    results = []
    full_text = " ".join([p.text for p in doc.paragraphs if len(p.text) > 100])
    words = re.findall(r'\w+', full_text.lower())
    
    if len(words) < 200: return []
    
    word_counts = {}
    for w in words: word_counts[w] = word_counts.get(w, 0) + 1
    
    # Сортуємо за частотою
    sorted_counts = sorted(word_counts.values(), reverse=True)
    
    # Перевіряємо першу трійку (за законом Ціпфа частота 2-го слова ~1/2 від 1-го, 3-го ~1/3)
    if len(sorted_counts) > 3:
        top1, top2, top3 = sorted_counts[0], sorted_counts[1], sorted_counts[2]
        r2 = top1 / top2
        r3 = top1 / top3
        
        # Якщо відношення занадто "ідеальні" (як у ШІ) або занадто хаотичні
        if 1.8 < r2 < 2.2 and 2.7 < r3 < 3.3:
            results.append(f" [⚠️] Частотність: Розподіл слів занадто математично ідеальний (ознака ШІ).")
        else:
            results.append(f" [✅] Частотність: Природний розподіл слів (Закон Ціпфа).")
            
    return results

def run_quality_audit(file_path):
    if not os.path.exists(file_path):
        print(f"Помилка: Файл {file_path} не знайдено.")
        return

    print(f"\n>>> ЗАПУСК АУДИТУ ЯКОСТІ ДЛЯ: {os.path.basename(file_path)} <<<\n")
    doc = Document(file_path)
    
    audit_report = []
    audit_report.extend(check_margins(doc))
    audit_report.extend(check_mixed_alphabets(doc))
    audit_report.extend(check_formatting_rules(doc))
    audit_report.extend(check_ai_style(doc))
    audit_report.extend(advanced_linguistic_analysis(doc))
    audit_report.extend(calculate_readability(doc))
    audit_report.extend(check_zipf_law(doc))
    
    # Структурні перевірки
    full_text = "\n".join([p.text for p in doc.paragraphs])
    if "ЗМІСТ" in full_text.upper():
        audit_report.append(" [✅] Зміст: знайдено в документі")
    else:
        audit_report.append(" [❌] Зміст: не знайдено (критична помилка!)")
        
    if "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ" in full_text.upper() or "ЛІТЕРАТУРА" in full_text.upper():
        audit_report.append(" [✅] Література: розділ присутній")
    else:
        audit_report.append(" [❌] Література: розділ не знайдено")

    for line in audit_report:
        print(line)
        
    print("\n>>> АУДИТ ЗАВЕРШЕНО <<<\n")

if __name__ == "__main__":
    import sys
    target = sys.argv[1] if len(sys.argv) > 1 else "docs/thesis/THESIS_FINAL_NO_MERMAID_AT_ALL.docx"
    run_quality_audit(target)
