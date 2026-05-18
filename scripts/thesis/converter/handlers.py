"""
БІБЛІОТЕКА ОБРОБНИКІВ СТРУКТУРНИХ ЕЛЕМЕНТІВ (Document Element Handlers)
======================================================================
Модуль реалізує логіку трансформації Markdown-компонентів у Word-об'єкти.
Забезпечує:
1. Dynamic Heading Processing: управління ієрархією заголовків (H1-H4) з фільтрацією для Змісту.
2. Smart Image Scaling: пропорційне масштабування зображень та підтримка кастомної ширини.
3. List & Table Orchestration: формування списків та таблиць (включаючи безрамкові бланки).
4. Logic Filtering: виключення службового контенту та "сміття" з академічного звіту.
5. Code Formatting: оформлення лістингів коду з використанням моноширинного шрифту 10pt.
Виступає виконавчим механізмом для побудови структури фінальної дипломної роботи.
"""
import os
import re
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from .styles import set_run_font, para_std, clean_inline, add_formatted_run

def clean_heading_dots(text):
    # Вимога керівника: "1.1 Назва" -> "1.1. Назва" (крапка в кінці номера)
    # Шукаємо номер на початку: цифри з крапками, за якими НЕ йде крапка
    return re.sub(r'^(\d+(?:\.\d+)*)(?!\.)(\s+)', r'\1.\2', text)

def should_be_in_toc(text, level=1):
    text_up = text.upper().strip()
    
    # 1. Специфічні виключення (сміття)
    trash_keywords = [
        "НА ТЕМУ", "ТЕМА:", "ЗАКЛАД ВИЩОЇ", "МІЖНАРОДНИЙ НАУКОВО-ТЕХНІЧНИЙ", 
        "БАКАЛАВР", "СПЕЦІАЛЬНІСТЮ", "КИЇВ", "СТУДЕНТУ", "ЛИТВИНЕНКУ", 
        "ОСВІТНЬОГО СТУПЕНЯ", "ВІДГУК", "РЕЦЕНЗІЯ", "РЕЦЕНЗЕНТА", "КЕРІВНИКА",
        "КАЛЕНДАРНИЙ ПЛАН", "КВАЛІФІКАЦІЙНА РОБОТА"
    ]
    if any(tk in text_up for tk in trash_keywords):
        return False

    # 2. Виключення дублів Реферату (якщо це H2)
    if level > 1 and text_up == "РЕФЕРАТ":
        return False
        
    # 3. Обов'язкові елементи (H1)
    must_toc_h1 = ["РЕФЕРАТ / ABSTRACT", "ВСТУП", "ВИСНОВКИ", "ЛІТЕРАТУРА", "ДОДАТКИ", "СКОРОЧЕНЬ", "РОЗДІЛ", "ЗАВДАННЯ", "З А В Д А Н Н Я"]
    if level == 1:
        if any(m in text_up for m in must_toc_h1):
            return True
        # Якщо це будь-який інший H1 (не сміття), додаємо
        return True

    # 4. Підрозділи (H2, H3) - додаємо тільки якщо є нумерація (1.1, А.1) або це Додаток
    if re.match(r'^(\d+\.|[А-Я]\.)', text) or "ДОДАТОК" in text_up:
        return True
        
    return False

def add_h1(doc, text, force_no_break=False):
    text_up = text.upper().strip()
    is_toc = should_be_in_toc(text, level=1)
    
    if is_toc:
        try: p = doc.add_paragraph(style='Heading 1')
        except: p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        p.paragraph_format.widow_control = True
    
    # Визначаємо, чи це головний розділ, який має починатися з нової сторінки
    major_keywords = ["РОЗДІЛ", "ВСТУП", "ВИСНОВКИ", "ЛІТЕРАТУРА", "ДОДАТКИ", "РЕФЕРАТ", "СКОРОЧЕНЬ", "ЗМІСТ", "ЗАВДАННЯ", "СПИСОК"]
    is_major = any(kw in text_up for kw in major_keywords)
    
    # Ставимо розрив сторінки ТІЛЬКИ для головних розділів і якщо не було force_no_break
    # ВИНЯТОК: "Київ" не повинен починатися з нової сторінки
    if is_major and not force_no_break and len(text_up) < 100 and "КИЇВ" not in text_up:
        p.paragraph_format.page_break_before = True
    else:
        p.paragraph_format.page_break_before = False
    
    # Вимога керівника: "РОЗДІЛ 1." (з крапкою)
    text = re.sub(r'^(РОЗДІЛ\s+\d+)(?!\.)', r'\1.', text, flags=re.IGNORECASE)
    
    # Центрування заголовків
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Універсальний парзер зі стилем H1 (16pt, AllCaps, Bold)
    add_formatted_run(p, text.upper(), size=16, bold_base=True)

def add_h2(doc, text):
    text = clean_heading_dots(text)
    is_toc = should_be_in_toc(text, level=2)
    if is_toc:
        try: p = doc.add_paragraph(style='Heading 2')
        except: p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()

    # Вимога керівника: H2 — жирний, по ширині
    add_formatted_run(p, text, size=14, bold_base=True)
    
    # Якщо це не нумерований підрозділ (на титулці), то центруємо
    if not re.match(r'^(\d+\.|[А-Я]\.)', text):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_h3(doc, text):
    text = clean_heading_dots(text)
    is_toc = should_be_in_toc(text, level=3)
    if is_toc:
        try: p = doc.add_paragraph(style='Heading 3')
        except: p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()

    # Вимога керівника: H3 — без курсиву, жирний, 14pt
    add_formatted_run(p, text, size=14, bold_base=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    pf.keep_together = True
    pf.widow_control = True

def add_h4(doc, text):
    # Вимога керівника: без курсиву
    text = clean_heading_dots(text)
    p = doc.add_paragraph()
    add_formatted_run(p, text, size=14, bold_base=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    pf.keep_together = True
    pf.widow_control = True

def add_body(doc, text, indent=True):
    if not text.strip(): return
    # Підтримка <br> для ручних розривів на титулці
    if text.strip() == "<br>":
        doc.add_paragraph().add_run("\u00A0")
        return
        
    p = doc.add_paragraph()
    # Вимога керівника: 14pt TNR, 1.5 інтервал, відступ 1.25, без зайвих відступів після
    add_formatted_run(p, text, size=14)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.keep_with_next = False
    pf.keep_together = False
    pf.widow_control = True
    if indent:
        pf.first_line_indent = Cm(1.25)
    else:
        pf.first_line_indent = Cm(0)

def add_list_item(doc, text, numbered=False):
    if not text.strip(): return
    if not numbered:
        try: p = doc.add_paragraph(style='List Bullet')
        except: p = doc.add_paragraph("• ")
    else:
        p = doc.add_paragraph()
    add_formatted_run(p, text, size=14)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5 # Міжрядковий інтервал 1.5
    pf.left_indent = Cm(0)
    pf.keep_with_next = False
    pf.keep_together = False
    pf.widow_control = True
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(2)

def add_image(doc, img_name, caption=None):
    target_width = None
    if "?" in img_name:
        parts = img_name.split("?")
        img_name = parts[0]
        if "w=" in parts[1]:
            try: target_width = float(parts[1].split("w=")[1].split("&")[0])
            except: pass
    img_path = os.path.join("docs", "images", img_name)
    if not os.path.exists(img_path): return

    # Вимога керівника: Enter перед рисунком
    doc.add_paragraph().add_run("\u00A0")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
            aspect = h_px / w_px
        if target_width: final_width = Cm(target_width)
        elif aspect > 1.2: final_width = Cm(8)
        elif aspect > 0.8: final_width = Cm(11)
        else: final_width = Cm(15.5)
        run.add_picture(img_path, width=final_width)
    except Exception as e: print(f" [ERR] Image ERR {img_name}: {e}")

def add_figure_caption(doc, text):
    # Вимога керівника: підпис рисунку відцентрований, курсивом, 12pt
    cap_clean = re.sub(r'^\*(.+)\*$', r'\1', text.strip())
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(0)
    run_cap = p_cap.add_run(cap_clean)
    set_run_font(run_cap, size=14, italic=False)
    
    # Вимога керівника: Джерело під рисунком
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_src.paragraph_format.space_before = Pt(0)
    p_src.paragraph_format.space_after = Pt(0)
    run_src = p_src.add_run("Джерело: згенеровано автором на основі програмного коду.")
    set_run_font(run_src, size=12, italic=False)
    
    # Enter після підпису
    doc.add_paragraph().add_run("\u00A0")

def add_code(doc, code_lines):
    # Вимога керівника: код в Courier New 12pt
    # Додатково: Автоматична підсвітка синтаксису (Keywords, Comments, Strings)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = Pt(16)
    
    python_keywords = {
        'def', 'class', 'import', 'from', 'return', 'if', 'else', 'elif', 
        'for', 'while', 'try', 'except', 'finally', 'with', 'as', 'in', 
        'is', 'and', 'or', 'not', 'None', 'True', 'False', 'self', 'lambda',
        'async', 'await', 'pass', 'break', 'continue', 'yield'
    }
    python_builtins = {
        'print', 'len', 'range', 'int', 'str', 'list', 'dict', 'set', 'tuple',
        'float', 'bool', 'object', 'super', 'type', 'enumerate', 'zip', 'sum',
        'min', 'max', 'abs', 'round', 'open', 'any', 'all', 'map', 'filter',
        'Exception', 'ValueError', 'TypeError', 'KeyError', 'IndexError', 'RuntimeError'
    }
    python_operators = {'=', '==', '!=', '<', '>', '<=', '>=', '+', '-', '*', '/', '//', '%', '**', '+=', '-=', '*=', '/='}

    # Простий розбір по рядках для підсвітки
    for line in code_lines:
        # Обробка коментарів (все після # - зелене)
        if '#' in line:
            parts = line.split('#', 1)
            main_part = parts[0]
            comment_part = '#' + parts[1]
        else:
            main_part = line
            comment_part = ""

        # Розбиваємо головну частину на токени
        tokens = re.split(r'(\s+|[(),\[\]{}:.=#@]|".*?"|\'.*?\')', main_part)
        for token in tokens:
            if not token: continue
            run = p.add_run(token)
            set_run_font(run, 11, mono=True)
            
            stripped = token.strip()
            # Ключові слова (Blue)
            if stripped in python_keywords:
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.bold = True
            # Вбудовані функції (Teal)
            elif stripped in python_builtins:
                run.font.color.rgb = RGBColor(0, 128, 128)
            # Рядки (Dark Red)
            elif (token.startswith('"') or token.startswith("'")):
                run.font.color.rgb = RGBColor(163, 21, 21)
            # Декоратори (Gold)
            elif token.startswith('@'):
                run.font.color.rgb = RGBColor(121, 94, 38)
            # Оператори (Slate Blue)
            elif stripped in python_operators:
                run.font.color.rgb = RGBColor(59, 89, 152)
            # Числа (Emerald)
            elif re.match(r'^\d+$', stripped):
                run.font.color.rgb = RGBColor(9, 134, 88)

        if comment_part:
            run = p.add_run(comment_part)
            set_run_font(run, 11, mono=True)
            run.font.color.rgb = RGBColor(0, 128, 0) # Green for comments
        
        # Перехід на новий рядок всередині параграфа (Shift+Enter)
        if line != code_lines[-1]:
            p.add_run('\n')

def add_table_caption(doc, text):
    # Вимога керівника: Таблиця 1.1 – Назва (коротке тире, по центру, без крапки в кінці)
    cap_clean = re.sub(r'^\*(.+)\*$', r'\1', text.strip())
    # Заміна крапки після номера на коротке тире
    cap_clean = re.sub(r'^(Таблиця\s+\d+\.\d+)\.?\s*', r'\1 – ', cap_clean)
    # Прибираємо крапку в кінці
    cap_clean = cap_clean.rstrip('.')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Таблиці по центру
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(cap_clean)
    set_run_font(run, size=14, italic=False) # Не курсив для таблиць за ДСТУ

def add_table(doc, table_lines):
    rows_data = []
    is_borderless = False
    for line in table_lines:
        if "<!-- NO_BORDER -->" in line:
            is_borderless = True
            line = line.replace("<!-- NO_BORDER -->", "")
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells if c): continue
        rows_data.append(cells)
    if not rows_data: return
    
    ncols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.allow_autofit = False # Вимикаємо автопідбір для повного контролю ширини
    
    if is_borderless:
        tbl.style = None
    else:
        tbl.style = 'Table Grid'

    # Робимо таблицю на всю ширину сторінки (16.5 см)
    available_width = 16.5
    col_widths = [0.0] * ncols
    
    # Спеціальна логіка для Щоденника (4 колонки: № | Зміст | Дати | Відмітка)
    if ncols == 4 and not is_borderless:
        col_widths = [1.2, 9.3, 3.5, 2.5]
    # Спеціальна логіка для Титулки (3 колонки: Виконав | Підпис | Прізвище)
    elif ncols == 3 and is_borderless:
        col_widths = [7.5, 4.0, 5.0] # Збільшуємо першу колонку для "Виконав (ла)..."
    else:
        # Пропорційно до тексту
        col_widths_chars = [0] * ncols
        for row_data in rows_data:
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < ncols:
                    col_widths_chars[c_idx] = max(col_widths_chars[c_idx], len(cell_text))
        
        total_chars = sum(col_widths_chars)
        if total_chars > 0:
            for c_idx in range(ncols):
                w = (col_widths_chars[c_idx] / total_chars) * available_width
                col_widths[c_idx] = max(w, 2.0)
            
            # Масштабуємо щоб в сумі було рівно 16.5
            curr_sum = sum(col_widths)
            col_widths = [w * (available_width / curr_sum) for w in col_widths]
        else:
            col_widths = [available_width / ncols] * ncols

    # Застосовуємо ширину колонок до об'єкта таблиці
    for c_idx in range(ncols):
        tbl.columns[c_idx].width = Cm(col_widths[c_idx])

    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0 and not is_borderless)
        row_data += [''] * (ncols - len(row_data))
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            # Явно встановлюємо ширину для кожної комірки (це надійніше ніж для всієї колонки)
            cell.width = Cm(col_widths[c_idx])
            
            cell.text = ""
            p = cell.paragraphs[0]
            # Для титулки (borderless) використовуємо 14pt, для звичайних таблиць 12pt
            font_size = 14 if is_borderless else 12
            add_formatted_run(p, cell_text, size=font_size, bold_base=is_header)
            for para in cell.paragraphs:
                pf = para.paragraph_format
                if is_borderless:
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER # Дані в таблицях зазвичай по центру
                # Налаштування відступів всередині комірок
                pf.space_before, pf.space_after = Pt(2), Pt(2)
                pf.line_spacing = 1.0
                
    # Enter після таблиці
    doc.add_paragraph().add_run("\u00A0")
