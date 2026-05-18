"""
ДВИГУН МАТЕМАТИЧНОЇ ТРАНСФОРМАЦІЇ (LaTeX-to-Word Math Engine)
===========================================================
Спеціалізований модуль для конвертації наукової нотації в нативні об'єкти Word.
Забезпечує:
1. LaTeX Parsing: вилучення та очищення математичних виразів з Markdown-контенту.
2. MathML Generation: перетворення формул у XML-представлення (MathML).
3. COM Automation Integration: вставлення формул у Word через буфер обміну (MS Word API).
4. Final Document Polishing: оновлення Змісту та полів документа після вставки об'єктів.
Гарантує професійний вигляд математичних доказів у дипломній роботі.
"""
import re
import os
import time
import shutil
import win32com
import win32com.client as win32
from win32com.client import dynamic as win32_dynamic
import win32clipboard
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from latex2mathml.converter import convert as latex_to_mathml
from .styles import set_run_font

from docx.enum.text import WD_TAB_ALIGNMENT

def _fix_word_com_cache():
    """Автоматично видаляє зіпсований win32com gen_py кеш Word для перегенерації. 
    Вирішує помилку: AttributeError 'CLSIDToPackageMap'."""
    try:
        gen_py_dir = os.path.join(os.path.dirname(win32com.__file__), 'gen_py')
        word_cache_id = '00020905-0000-0000-C000-000000000046x0x8x7'
        cache_path = os.path.join(gen_py_dir, word_cache_id)
        if os.path.exists(cache_path):
            shutil.rmtree(cache_path, ignore_errors=True)
            print(" [INFO] win32com Word cache видалено (був зіпсований), перегенерується...")
    except Exception as e:
        print(f" [WARN] Cache clear: {e}")

def add_formula(doc, latex_text, chapter_num, formula_num, has_where=False):
    latex = re.sub(r'^\$\$?|\$\$?$', '', latex_text.strip()).strip()
    # Видаляємо ручну нумерацію, якщо вона вже була (щоб уникнути дублювання)
    latex = re.sub(r'\\quad\s*\(\d+\.\d+\)$', '', latex).strip()
    latex = re.sub(r'\(\d+\.\d+\)$', '', latex).strip()
    if not latex: return
    
    # Якщо наступний рядок починається з "де", додаємо кому до формули
    if has_where and not latex.endswith(','):
        latex += ','
    
    # Enter перед формулою
    doc.add_paragraph().add_run("\u00A0")
    
    p = doc.add_paragraph()
    
    # Налаштування табуляції: центр (8.5 см) та правий край (17.0 см)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(8.5), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
    
    # Формат: <Tab> Формула <Tab> (Розділ.Номер)
    run = p.add_run(f"\t⇲{latex}⇱\t({chapter_num}.{formula_num})")
    set_run_font(run, 14)
    
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Enter після формули
    doc.add_paragraph().add_run("\u00A0")

def paste_mathml(text):
    for _ in range(10):
        try:
            win32clipboard.OpenClipboard()
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardText(text, win32clipboard.CF_UNICODETEXT)
            win32clipboard.CloseClipboard()
            return
        except Exception: time.sleep(0.1)
    raise Exception("Clipboard Locked!")

def finalize_thesis_document(docx_path):
    # 1. Видаляємо зіпсований gen_py файл
    _fix_word_com_cache()
    # 2. Monkey-patch gencache: повертає None для всіх CLSID —
    #    це змушує __WrapDispatch використовувати звичайний CDispatch (пізнє зв’язування)
    try:
        import win32com.client.gencache as _gc
        _gc.GetClassForCLSID = lambda clsid: None
    except:
        pass
    print("Запуск фінальної оптимізації документа (формули + зміст)...")
    
    import subprocess
    subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE", "/T"], capture_output=True)
    time.sleep(1)
    
    # dynamic.Dispatch + monkey-patched gencache = ніяких посилань на кеш
    try:
        word = win32_dynamic.Dispatch("Word.Application")
    except Exception as e:
        print(f" [ERR] Word COM недоступний, формули пропускаються: {e}")
        return
        
    word.Visible = False
    word.DisplayAlerts = 0
    
    try:
        abs_path = os.path.abspath(docx_path)
        doc_word = word.Documents.Open(abs_path)
        
        # 1. ПЕРЕТВОРЕННЯ ФОРМУЛ
        word.Selection.HomeKey(Unit=6)
        find = word.Selection.Find
        find.ClearFormatting()
        math_count = 0
        
        while True:
            find.Text = "⇲"
            if not find.Execute(): break
            start_pos = word.Selection.Start
            word.Selection.Collapse(Direction=0)
            find.Text = "⇱"
            if not find.Execute(): break
            end_pos = word.Selection.End
            rng = doc_word.Range(start_pos, end_pos)
            formula = rng.Text.replace("⇲", "").replace("⇱", "").strip().replace(r'\quad', r'\ ')
            try:
                mathml = latex_to_mathml(formula)
                paste_mathml(mathml)
                time.sleep(0.05)
                rng.Paste()
                math_count += 1
            except: pass
            word.Selection.Collapse(Direction=0)
        
        # 2. ОНОВЛЕННЯ ЗМІСТУ ТА ПОЛІВ
        print(f"✅ Формули ({math_count}) оброблено. Оновлення Змісту...")
        doc_word.Fields.Update()
        for toc in doc_word.TablesOfContents:
            toc.Update()
        
        doc_word.Save()
        print("✅ Документ повністю оптимізовано за один прохід.")
    finally:
        try:
            doc_word.Close()
            word.Quit()
        except: pass
