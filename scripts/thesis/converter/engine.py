"""
ДВИГУН ФОРМАТУВАННЯ ТА КОНВЕРТАЦІЇ ТЕЗИ (Thesis Layout Engine)
=============================================================
Ядро системи трансформації Markdown-тексту в професійно оформлений Word-документ.
Ключові можливості:
1. Intelligent Markdown Parsing: розпізнавання заголовків, списків, зображень та блоків коду.
2. Word Object Transformation: перетворення LaTeX-формул у нативні об'єкти Word.
3. Academic Formatting Enforcement: налаштування відступів та шрифтів за вимогами ДСТУ.
4. Dynamic TOC & Pagination: автоматична генерація змісту та нумерації сторінок.
5. Mermaid Filtering: інтелектуальне очищення схем, що призначені лише для веб-версії.
Забезпечує високу якість фінальної верстки пояснювальної записки до диплома.
"""
import os
import re
import time
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .config import MERMAID_MAP
from .styles import add_page_numbers, add_formatted_run
from .handlers import (
    add_h1, add_h2, add_h3, add_h4, add_body,
    add_list_item, add_image, add_figure_caption, add_table_caption, add_code, add_table
)
from .formulas import add_formula, finalize_thesis_document

NAV_PATTERN = re.compile(r'\[.*?(Назад до|Далі:|Повернутись|⬅️|➡️).*?\]\(.*?\.md\)')

def run_conversion(input_md, output_docx, include_appendix=False):
    if not os.path.exists(input_md):
        print(f" [ERR] Input not found: {input_md}")
        return

    with open(input_md, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    
    # СТВОРЕННЯ ГЛОБАЛЬНОГО СТИЛЮ (Звичайна верстка)
    style_norm = doc.styles['Normal']
    style_norm.font.name = 'Times New Roman'
    style_norm.font.size = Pt(14)
    pf_norm = style_norm.paragraph_format
    pf_norm.space_after = Pt(0)
    pf_norm.line_spacing = 1.5
    # ГАЛОЧКИ ДЛЯ ТЕКСТУ: Тільки контроль висячих рядків!
    pf_norm.keep_with_next = False
    pf_norm.keep_together = False
    pf_norm.widow_control = True

    # СТВОРЕННЯ СТИЛЮ ДЛЯ ЗАГОЛОВКІВ (Heading 1)
    if 'Heading 1' not in doc.styles:
        doc.styles.add_style('Heading 1', 1)
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 4-та ГАЛОЧКА (Page Break Before) - ВИМКНЕНА ЗА ЗАМОВЧУВАННЯМ (для Титулки)
    h1_style.paragraph_format.page_break_before = False
    h1_style.paragraph_format.keep_with_next = True
    h1_style.paragraph_format.keep_together = True
    h1_style.paragraph_format.widow_control = True

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin, section.right_margin = Cm(3.0), Cm(1.0)
    section.top_margin, section.bottom_margin = Cm(2.0), Cm(2.0)

    in_code = False
    in_math = False
    in_table = False
    in_mermaid = False
    code_buf = []
    math_buf = []
    table_buf = []
    last_was_image = False
    
    current_chapter = 0
    current_formula = 0
    in_mermaid = False
    just_had_pagebreak = True # Ставимо True на початку, щоб перший заголовок не робив зайвий розрив

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        i += 1
        
        stripped = line.strip()

        # Ігноруємо блоки Mermaid та їхні підписи (вони тільки для GitHub)
        if stripped.startswith('```mermaid'):
            in_mermaid = True
            continue
        if in_mermaid:
            if stripped.startswith('```'):
                in_mermaid = False
            continue
            
        # Ігноруємо підписи до Mermaid схем
        if re.search(r'Схема\s+\d+\.\d+\..*\(Mermaid-версія', stripped):
            continue

        # Фільтрація навігації
        if NAV_PATTERN.search(line) and len(line.strip()) < 150: continue

        # Блоки коду
        if line.startswith('```'):
            last_was_image = False
            if in_code:
                if code_lang != "mermaid":
                    add_code(doc, code_buf)
                code_buf = []; in_code = False; code_lang = ""
            else:
                in_code = True
                code_lang = line[3:].strip().lower()
            continue
        if in_code: code_buf.append(line); continue

        # Блоки формул $$
        if line.strip().startswith('$$'):
            last_was_image = False
            if in_math:
                full_math = " ".join(math_buf) + " " + line.strip().replace('$$', '')
                current_formula += 1
                ch_num = current_chapter if current_chapter > 0 else 1
                
                # Перевіряємо наступний рядок на "де " (для коми)
                has_where = False
                if i < len(lines):
                    next_l = lines[i].strip()
                    if next_l.startswith("де ") or next_l == "де":
                        has_where = True
                
                add_formula(doc, full_math.strip(), ch_num, current_formula, has_where=has_where)
                math_buf = []; in_math = False
            else:
                if line.strip().endswith('$$') and len(line.strip()) > 2:
                    current_formula += 1
                    ch_num = current_chapter if current_chapter > 0 else 1
                    
                    # Перевіряємо наступний рядок на "де "
                    has_where = False
                    if i < len(lines):
                        next_l = lines[i].strip()
                        if next_l.startswith("де ") or next_l == "де":
                            has_where = True
                            
                    add_formula(doc, line.strip(), ch_num, current_formula, has_where=has_where)
                else:
                    in_math = True
                    math_buf = [line.strip().replace('$$', '')]
            continue
        if in_math: math_buf.append(line.strip()); continue

        # Таблиці
        if line.strip().startswith('|'):
            last_was_image = False
            if not in_table: in_table, table_buf = True, []
            table_buf.append(line); continue
        elif in_table:
            add_table(doc, table_buf); table_buf = []; in_table = False

        if not stripped or re.match(r'^---+$', stripped): continue

        # Обробка інлайнових формул у тексті
        def process_inline_math(text):
            return re.sub(r'\$(.+?)\$', r'⇲\1⇱', text)

        # Обробка підпису під рисунком
        if re.search(r'^\*Рис\..+\*', stripped):
            if not last_was_image:
                p = doc.add_paragraph()
                run = p.add_run("[УВАГА: ВІДСУТНЄ ЗОБРАЖЕННЯ ДЛЯ ЦЬОГО ПІДПИСУ!]")
                run.font.bold = True
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(255, 0, 0)
            
            add_figure_caption(doc, stripped)
            last_was_image = False # Скидаємо після підпису
            continue

        # Обробка підпису над таблицею
        if re.search(r'^\*Таблиця\s+\d+\.\d+.*\*', stripped):
            add_table_caption(doc, stripped)
            continue

        if stripped.startswith('####'): 
            text = re.sub(r'^####\s*', '', stripped)
            add_h4(doc, process_inline_math(text))
            last_was_image = False
        elif stripped.startswith('###'): 
            text = re.sub(r'^###\s*', '', stripped)
            add_h3(doc, process_inline_math(text))
            last_was_image = False
        elif stripped.startswith('##'): 
            text = re.sub(r'^##\s*', '', stripped)
            add_h2(doc, process_inline_math(text))
            last_was_image = False
        elif stripped == "<pagebreak>":
            doc.add_page_break()
            just_had_pagebreak = True
            continue
        elif stripped.startswith('<p align="center">'):
            # Обробка центрування (Додатки та бланки)
            clean_text = re.sub(r'<.*?>', '', stripped) # Видаляємо всі теги
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = "<b>" in stripped or "**" in stripped
            add_formatted_run(p, clean_text, size=14, bold_base=is_bold)
            continue
            
        elif stripped.startswith('#'):
            text = re.sub(r'^#\s*', '', stripped)
            # ВСТАВКА ЗМІСТУ ПЕРЕД СКОРОЧЕННЯМИ (Згідно п. 6.1: Реферат -> Зміст -> Скорочення)
            if "СКОРОЧЕНЬ" in text.upper() or "ПОЗНАЧЕНЬ" in text.upper():
                p_toc = doc.add_paragraph()
                p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_formatted_run(p_toc, "ЗМІСТ", size=16, bold_base=True)
                
                p_field = doc.add_paragraph()
                run = p_field.add_run()
                # Вставка поля TOC (Table of Contents)
                fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText'); instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
                fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3)
                # Після змісту теж ставимо прапорець, щоб наступний заголовок не стрибав
                just_had_pagebreak = True
                # ПРИБРАНО ПОВТОРНИЙ doc.add_page_break(), бо add_h1 сам зробить розрив для наступного розділу

            if text.upper().startswith('РОЗДІЛ'):
                m = re.search(r'РОЗДІЛ\s+(\d+)', text.upper())
                if m:
                    current_chapter = int(m.group(1))
                    current_formula = 0

            add_h1(doc, process_inline_math(text), force_no_break=just_had_pagebreak)
            just_had_pagebreak = False
            last_was_image = False
        elif re.match(r'^Додаток\s+[А-Яа-яA-Za-z]', stripped): add_h1(doc, stripped); last_was_image = False
        elif re.search(r'!\[(.*?)\]\((.*?)\)', stripped):
            m = re.search(r'!\[(.*?)\]\((.*?)\)', stripped)
            add_image(doc, os.path.basename(m.group(2)))
            last_was_image = True
        elif re.match(r'^(\*\*)?\d+\.\s', stripped): 
            add_list_item(doc, process_inline_math(stripped), numbered=True)
            last_was_image = False
        elif stripped.startswith('* ') or stripped.startswith('- '): 
            add_list_item(doc, process_inline_math(stripped[2:]))
            last_was_image = False
        else: 
            # Вимога керівника: "де" без відступу
            if stripped.startswith("де ") or stripped == "де":
                add_body(doc, process_inline_math(line), indent=False)
            else:
                add_body(doc, process_inline_math(line))
            
            if stripped: 
                last_was_image = False
                just_had_pagebreak = False


    if in_table: add_table(doc, table_buf)
    
        
    add_page_numbers(doc)
    doc.save(output_docx)
    print(f"✅ Документ збережено: {output_docx}")
    
    # ЕТАП ОПТИМІЗАЦІЇ (ОДИН ПРОХІД WORD)
    finalize_thesis_document(output_docx)

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python engine.py input.md output.docx")
    else:
        run_conversion(sys.argv[1], sys.argv[2])
