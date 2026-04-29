import os
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from .styles import set_run_font, para_std, clean_inline, add_formatted_run

def clean_heading_dots(text):
    # Прибирає крапку після номерів підрозділів (наприклад "1.1. Назва" -> "1.1 Назва")
    # Але не чіпає "РОЗДІЛ 1." або інші випадки
    return re.sub(r'^(\d+(\.\d+)+)\.\s+', r'\1 ', text)

def should_be_in_toc(text, level=1):
    text_up = text.upper().strip()
    
    # 1. Специфічні виключення (сміття)
    trash_keywords = [
        "НА ТЕМУ:", "ТЕМА:", "ЗАКЛАД ВИЩОЇ", "МІЖНАРОДНИЙ НАУКОВО-ТЕХНІЧНИЙ", 
        "БАКАЛАВР", "СПЕЦІАЛЬНІСТЮ", "КИЇВ –", "СТУДЕНТУ", "ЛИТВИНЕНКУ", 
        "ОСВІТНЬОГО СТУПЕНЯ", "ВІДГУК", "РЕЦЕНЗІЯ", "РЕЦЕНЗЕНТА", "КЕРІВНИКА",
        "КАЛЕНДАРНИЙ ПЛАН"
    ]
    if any(tk in text_up for tk in trash_keywords):
        return False

    # 2. Виключення дублів Реферату (якщо це H2)
    if level > 1 and text_up == "РЕФЕРАТ":
        return False
        
    # 3. Обов'язкові елементи (H1)
    must_toc_h1 = ["РЕФЕРАТ / ABSTRACT", "ВСТУП", "ВИСНОВКИ", "ЛІТЕРАТУРА", "ДОДАТКИ", "СКОРОЧЕНЬ", "РОЗДІЛ", "ЗАВДАННЯ", "З А В Д А Н Н Я"]
    if level == 1:
        if any(m in text_up for m in must_toc_h1):
            return True
        # Якщо це будь-який інший H1 (не сміття), додаємо
        return True

    # 4. Підрозділи (H2, H3) - додаємо тільки якщо є нумерація (1.1, А.1) або це Додаток
    if re.match(r'^(\d+\.|[А-Я]\.)', text) or "ДОДАТОК" in text_up:
        return True
        
    return False

def add_h1(doc, text):
    text_up = text.upper().strip()
    is_toc = should_be_in_toc(text, level=1)
    if is_toc:
        try: p = doc.add_paragraph(style='Heading 1')
        except: p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        
    # Прибираємо крапку після "РОЗДІЛ 1."
    text = re.sub(r'^(РОЗДІЛ\s+\d+)\.', r'\1', text, flags=re.IGNORECASE)
    
    # Універсальний парзер зі стилем H1 (AllCaps + Bold)
    add_formatted_run(p, text.upper(), size=16, bold_base=True)
    pf = p.paragraph_format
    pf.alignment          = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before       = Pt(12)
    pf.space_after        = Pt(6)
    pf.line_spacing       = 1.0 # Одинарний інтервал за п. 6.1
    pf.first_line_indent  = Cm(0)
    
    # Розрив сторінки: тільки для головних вузлів. ВІДГУК/РЕЦЕНЗІЯ йдуть після ЗАКЛАД, тому їм не треба розрив.
    break_keywords = ["РОЗДІЛ", "ЗАКЛАД", "ЗАВДАННЯ", "З А В Д А Н Н Я", "РЕФЕРАТ / ABSTRACT", "ВСТУП", "ВИСНОВКИ", "СПИСОК", "ДОДАТКИ"]
    if is_toc or any(bk in text_up for bk in break_keywords):
        pf.page_break_before = True

def add_h2(doc, text):
    text = clean_heading_dots(text)
    is_toc = should_be_in_toc(text, level=2)
    if is_toc:
        try: p = doc.add_paragraph(style='Heading 2')
        except: p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
    
    # Налаштування згідно п. 6.1: по ширині, відступ 1.25 см
    add_formatted_run(p, text, size=14, bold_base=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    pf.first_line_indent = Cm(1.25)

def add_h3(doc, text):
    text = clean_heading_dots(text)
    is_toc = should_be_in_toc(text, level=3)
    if is_toc:
        try: p = doc.add_paragraph(style='Heading 3')
        except: p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
    
    # Налаштування згідно п. 6.1: по ширині, відступ 1.25 см
    add_formatted_run(p, text, size=14, bold_base=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    pf.first_line_indent = Cm(1.25)

def add_h4(doc, text):
    text = clean_heading_dots(text)
    p = doc.add_paragraph()
    add_formatted_run(p, text, size=14, italic_base=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(2); pf.space_after = Pt(2)
    pf.line_spacing = Pt(18)
    pf.first_line_indent = Cm(1.25)

def add_body(doc, text):
    if not text.strip(): return
    p = doc.add_paragraph()
    add_formatted_run(p, text, size=14)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(4)
    pf.line_spacing = 1.5 # Міжрядковий інтервал 1.5
    pf.first_line_indent = Cm(1.25)

def add_list_item(doc, text, numbered=False):
    if not text.strip(): return
    if not numbered:
        try: p = doc.add_paragraph(style='List Bullet')
        except: p = doc.add_paragraph("• ")
    else:
        p = doc.add_paragraph()
    add_formatted_run(p, text, size=14)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = 1.5 # Міжрядковий інтервал 1.5
    pf.left_indent = Cm(1.25)
    if numbered: pf.first_line_indent = Cm(-0.75)
    pf.space_after = Pt(2)

def add_image(doc, img_name, caption):
    target_width = None
    if "?" in img_name:
        parts = img_name.split("?")
        img_name = parts[0]
        if "w=" in parts[1]:
            try: target_width = float(parts[1].split("w=")[1].split("&")[0])
            except: pass
    img_path = os.path.join("docs", "images", img_name)
    if not os.path.exists(img_path): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size
            aspect = h_px / w_px
        if target_width: final_width = Cm(target_width)
        elif aspect > 1.2: final_width = Cm(8)
        elif aspect > 0.8: final_width = Cm(11)
        else: final_width = Cm(15.5)
        run.add_picture(img_path, width=final_width)
    except Exception as e: print(f" [ERR] Image ERR {img_name}: {e}")

def add_code(doc, code_lines):
    p = doc.add_paragraph()
    run = p.add_run("\n".join(code_lines))
    set_run_font(run, 10, mono=True)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(3); pf.space_after = Pt(3)
    pf.line_spacing = Pt(14)

def add_table(doc, table_lines):
    rows_data = []
    is_borderless = False
    for line in table_lines:
        if "<!-- NO_BORDER -->" in line:
            is_borderless = True
            line = line.replace("<!-- NO_BORDER -->", "")
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells if c): continue
        rows_data.append(cells)
    if not rows_data: return
    ncols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    if is_borderless:
        tbl.style = None
    else:
        tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0 and not is_borderless)
        row_data += [''] * (ncols - len(row_data))
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # 11pt для офіційних таблиць в бланках
            add_formatted_run(p, cell_text, size=11, bold_base=is_header)
            for para in cell.paragraphs:
                pf = para.paragraph_format
                if is_borderless:
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Максимальне ущільнення для бланків
                pf.space_before, pf.space_after, pf.line_spacing = Pt(1), Pt(1), Pt(14)
