import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_run_font(run, size=14, bold=False, italic=False, mono=False):
    fname = "Courier New" if mono else "Times New Roman"
    run.font.name  = fname
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), fname)
    rFonts.set(qn("w:cs"),    fname)
    rPr.insert(0, rFonts)

def add_formatted_run(paragraph, text, size=14, bold_base=False, italic_base=False, mono=False):
    """Парсить текст з підтримкою **bold**, *italic* та <u>underline</u> і додає run-и в параграф."""
    text = text.replace("<br>", "\n").replace("<BR>", "\n")
    pattern = r'(\*\*|\*|<u>|</u>)'
    parts = re.split(pattern, text)
    curr_bold = bold_base
    curr_italic = italic_base
    curr_under = False
    for part in parts:
        if part == "**":
            curr_bold = not curr_bold
            continue
        elif part == "*":
            curr_italic = not curr_italic
            continue
        elif part == "<u>":
            curr_under = True
            continue
        elif part == "</u>":
            curr_under = False
            continue
        if part:
            run = paragraph.add_run(part)
            set_run_font(run, size, bold=curr_bold, italic=curr_italic, mono=mono)
            run.underline = curr_under

def clean_inline(text):
    """Прибирає Markdown розмітку, залишає формули з мітками."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\$\$(.+?)\$\$', r'⇲\1⇱', text, flags=re.DOTALL)
    text = re.sub(r'\$([^\$\n]+?)\$', r'⇲\1⇱', text)
    return text.strip()

def para_std(p, indent=True):
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(21)
    if indent:
        pf.first_line_indent = Cm(1.25)

def add_page_numbers(doc):
    section = doc.sections[0]
    header  = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    for tag, val in [('begin', None), ('instrText', 'PAGE'), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.text = val
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
