"""
Конвертація THESIS_FULL_FINAL_UTF8.md -> THESIS_FINAL.docx
З підтримкою: таблиць, формул (OMML), списків, коду, заголовків
"""
import re
import os
import win32com.client as win32
import win32clipboard
from latex2mathml.converter import convert as latex_to_mathml
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

INPUT     = "docs/thesis/THESIS_FULL_FINAL_UTF8.md"
OUTPUT    = "docs/thesis/THESIS_FINAL.docx"

with open(INPUT, encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()

# ── Поля сторінки (ДСТУ) ──────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(1.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── Утиліти шрифтів ───────────────────────────────────────────────
def set_run_font(run, size=14, bold=False, italic=False, mono=False):
    fname = "Courier New" if mono else "Times New Roman"
    run.font.name  = fname
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), fname)
    rFonts.set(qn("w:cs"),    fname)
    rPr.insert(0, rFonts)

def clean_inline(text):
    """Прибирає Markdown розмітку, залишає формули з мітками."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Очистка залишків HTML-тегів (зокрема </p>, <br> тощо)
    text = re.sub(r'<[^>]+>', '', text)
    # Зберігаємо формули з унікальними мітками для нежадібного пошуку у Word
    text = re.sub(r'\$\$(.+?)\$\$', r'⇲\1⇱', text, flags=re.DOTALL)
    text = re.sub(r'\$([^\$\n]+?)\$', r'⇲\1⇱', text)
    return text.strip()

def para_std(p, indent=True):
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(21)
    if indent:
        pf.first_line_indent = Cm(1.25)

# ── Типи абзаців ──────────────────────────────────────────────────
def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_run_font(run, 16, bold=True)
    pf = p.paragraph_format
    pf.alignment          = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before       = Pt(12)
    pf.space_after        = Pt(6)
    pf.line_spacing       = Pt(24)
    pf.first_line_indent  = Cm(0)
    pf.page_break_before  = True

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 14, bold=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    pf.line_spacing = Pt(21)
    pf.first_line_indent = Cm(1.25)

def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 14, bold=True, italic=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    pf.line_spacing = Pt(21)
    pf.first_line_indent = Cm(1.25)

def add_h4(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 14, italic=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(3); pf.space_after = Pt(3)
    pf.line_spacing = Pt(21)
    pf.first_line_indent = Cm(1.25)

def add_body(text):
    text = clean_inline(text)
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 14)
    para_std(p)

def add_formula(latex_text):
    """Генерація блокових формул з мітками."""
    latex = re.sub(r'^\$\$?|\$\$?$', '', latex_text.strip()).strip()
    if not latex:
        return
    p = doc.add_paragraph()
    run = p.add_run(f"⇲{latex}⇱")
    set_run_font(run, 14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)

def add_list_item(text, numbered=False):
    text = clean_inline(text)
    if not text:
        return
    # Для маркованого списку використовуємо стандартний стиль
    if not numbered:
        try:
            p = doc.add_paragraph(style='List Bullet')
        except:
            p = doc.add_paragraph("• ")
    else:
        p = doc.add_paragraph()
        
    run = p.add_run(text)
    set_run_font(run, 14)
    pf = p.paragraph_format
    # Для списків краще вирівнювання по лівому краю, щоб не було гігантських пробілів
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = Pt(21)
    
    if numbered:
        pf.left_indent = Cm(1.25)
        pf.first_line_indent = Cm(-0.75)
    else:
        pf.left_indent = Cm(1.25)

def add_image(img_name, caption):
    """Шукає фото в docs/images та вставляє в Word з урахуванням пропорцій."""
    # Обробка параметра ширини ![alt](file.png?w=10)
    target_width = None
    if "?" in img_name:
        parts = img_name.split("?")
        img_name = parts[0]
        params = parts[1]
        if "w=" in params:
            try:
                target_width = float(params.split("w=")[1].split("&")[0])
            except: pass

    img_path = os.path.join("docs", "images", img_name)
    if not os.path.exists(img_path):
        print(f" [WARN] Зображення не знайдено: {img_path}")
        return

    # Вставляємо картинку
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    
    try:
        # Аналізуємо розмір через PIL
        with Image.open(img_path) as img:
            w_px, h_px = img.size
            aspect = h_px / w_px
        
        # Логіка підбору розміру
        if target_width:
            final_width = Cm(target_width)
        elif aspect > 1.2: # Дуже вертикальне (як ваша діаграма)
             final_width = Cm(8)  # Робимо вужчим, щоб не було на всю сторінку
        elif aspect > 0.8: # Квадратне або злегка вертикальне
             final_width = Cm(11)
        else: # Горизонтальне (скріншоти інтерфейсу)
             final_width = Cm(15.5)

        run.add_picture(img_path, width=final_width)
    except Exception as e:
        print(f" [ERR] Не вдалося вставити {img_name}: {e}")
        return

    # Додаємо підпис під картинкою (ВИДАЛЕНО, бо підпис іде наступним рядком у MD)
    pass

def add_code(code_lines):
    """Блок коду: Courier New 10pt."""
    p = doc.add_paragraph()
    run = p.add_run("\n".join(code_lines))
    set_run_font(run, 10, mono=True)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(3); pf.space_after = Pt(3)
    pf.line_spacing = Pt(14)

# ── Парсинг таблиць ──────────────────────────────────────────────
def parse_md_table(table_lines):
    """Перетворює рядки Markdown-таблиці у list[list[str]]."""
    rows = []
    for line in table_lines:
        line = line.strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        # Пропускаємо роздільник :---:
        if all(re.match(r'^:?-+:?$', c) for c in cells if c):
            continue
        rows.append(cells)
    return rows

def add_table(table_lines):
    rows = parse_md_table(table_lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    # Вирівнюємо кількість колонок
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell_text = clean_inline(cell_text)
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, 12, bold=is_header)
                pf = para.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after  = Pt(2)
                pf.line_spacing = Pt(18)

# ── Нумерація сторінок (верхній правий кут згідно з п. 208 КАНБАН) ─────
def add_page_numbers():
    section = doc.sections[0]
    header  = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    for tag, val in [('begin', None), ('instrText', 'PAGE'), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.text = val
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

# ── Головний цикл обробки ─────────────────────────────────────────
in_code       = False
code_buf      = []
in_table      = False
table_buf     = []
table_buf     = []

i = 0
while i < len(lines):
    raw  = lines[i]
    line = raw.rstrip('\r\n')
    i += 1

    # ── Словник замін Mermaid на фото ────────────────────────────
    MERMAID_MAP = {
        "graph LR\n    subgraph RNN": ("diag_lstm_compare.png", "Рис. 1.1. Схематичне порівняння архітектур Simple RNN та LSTM"),
        "flowchart LR\n    User":    ("diag_use_case.png",    "Рис. 3.1. Діаграма прецедентів системи EnergyMonitor"),
        "stateDiagram-v2":            ("diag_sequence_1.png",  "Рис. 3.3. Діаграма послідовності (Частина 1)"),
        "graph TB\n    subgraph UI": ("diag_architecture.png", "Рис. 3.0. Архітектурна схема системи EnergyMonitor-OLAP"),
        "graph TD\n    subgraph Local": ("diag_infra_cloud.png", "Рис. 3.2. Схема розгортання та потоків даних системи"),
        "sequenceDiagram":            ("diag_sequence_horizontal.png", "Рис. 3.1. Детальна діаграма послідовності взаємодії компонентів"),
        "erDiagram":                  ("diag_er_db.png",       "Рис. 3.4. Схема бази даних (ER-діаграма)"),
        "flowchart LR\n    Push":    ("diag_cicd_pipeline.png", "Рис. 3.5. Схема CI/CD конвеєра автоматизації")
    }

    # ── Блок коду ``` ... ```
    if line.startswith('```'):
        if in_code:
            # Перевіряємо, чи це не Mermaid, який ми хочемо замінити на фото
            full_code = "\n".join(code_buf)
            replaced = False
            for key, (img, cap) in MERMAID_MAP.items():
                if key in full_code:
                    # Перевіряємо наявність файла
                    if os.path.exists(os.path.join("docs", "images", img)):
                        add_image(img, cap)
                        if img == "diag_sequence_1.png":
                            add_image("diag_sequence_2.png", "Рис. 3.3. Діаграма послідовності (Частина 2)")
                        replaced = True
                        break
            
            if not replaced:
                add_code(code_buf)
            
            code_buf = []; in_code = False
        else:
            in_code = True
        continue
    if in_code:
        code_buf.append(line)
        continue

    # ── Таблиця
    if line.strip().startswith('|'):
        if not in_table:
            in_table = True
            table_buf = []
        table_buf.append(line)
        continue
    else:
        if in_table:
            add_table(table_buf)
            table_buf = []; in_table = False

    stripped = line.strip()

    # Порожній рядок
    if not stripped:
        continue

    # Горизонтальна лінія
    if re.match(r'^---+$', stripped):
        continue

    # Одиночна формула $...$  або $$...$$  на окремому рядку
    if re.match(r'^\$\$.+\$\$$', stripped):
        add_formula(stripped)
        continue

    # Заголовки
    if stripped.startswith('#### '):
        add_h4(stripped[5:])
        continue
    elif stripped.startswith('### '):
        add_h3(stripped[4:])
        continue
    elif stripped.startswith('## '):
        add_h2(stripped[3:])
        continue
    elif stripped.startswith('# '):
        add_h1(stripped[2:])
        continue
    
    # Автоматичне оформлення додатків (Додаток А, Б...)
    if re.match(r'^Додаток\s+[А-Яа-яA-Za-z]', stripped):
        add_h1(stripped)
        continue

    # Зображення ![Alt](file.png)
    img_match = re.search(r'!\[(.*?)\]\((.*?)\)', stripped)
    if img_match:
        caption = img_match.group(1)
        img_file = os.path.basename(img_match.group(2))
        add_image(img_file, caption)
        continue

    # Нумерований список (розуміє і 3., і **3.)
    if re.match(r'^(\*\*)?\d+\.\s', stripped):
        add_list_item(stripped, numbered=True)
        continue
    # Маркований список
    elif stripped.startswith('* ') or stripped.startswith('- '):
        add_list_item(stripped[2:])

    # Звичайний текст
    else:
        add_body(line)

# Якщо файл закінчився всередині таблиці
if in_table:
    add_table(table_buf)

add_page_numbers()
doc.save(OUTPUT)
print(f"✅ Базовий документ збережено: {OUTPUT}")


print("Запуск COM-автоматизації для перетворення 2D формул...")

import time

def paste_mathml(text):
    for _ in range(10):
        try:
            win32clipboard.OpenClipboard()
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardText(text, win32clipboard.CF_UNICODETEXT)
            win32clipboard.CloseClipboard()
            return
        except Exception:
            time.sleep(0.1)
    raise Exception("Буфер обміну заблоковано!")

word = win32.Dispatch("Word.Application")
word.Visible = False

try:
    doc_word = word.Documents.Open(os.path.abspath(OUTPUT))
    word.Selection.HomeKey(Unit=6) # Початок документу
    
    find = word.Selection.Find
    find.ClearFormatting()
    find.MatchWildcards = False  # ВИМИКАЄМО Wildcards!
    
    math_count = 0
    while True:
        # Шукаємо старт
        find.Text = "⇲"
        if not find.Execute():
            break
        start_pos = word.Selection.Start
        
        # Шукаємо кінець
        word.Selection.Collapse(Direction=0)
        find.Text = "⇱"
        if not find.Execute():
            break
        end_pos = word.Selection.End
        
        # Виділяємо весь блок
        rng = doc_word.Range(start_pos, end_pos)
        formula = rng.Text.replace("⇲", "").replace("⇱", "").strip()
        formula = formula.replace(r'\quad', r'\ ')
        
        try:
            mathml = latex_to_mathml(formula)
            paste_mathml(mathml)
            time.sleep(0.05)   # Мікро-пауза, щоб Word встиг прочитати буфер
            rng.Paste()        # Пряма вставка в об'єкт без рухання курсору
            math_count += 1
        except Exception as e:
            # Запобіжник для "битих" формул
            print(f"Помилка [{formula}]: {e}")
            rng.Text = f"[{formula}]"
            
        word.Selection.Collapse(Direction=0)
        
    doc_word.Save()
    print(f"✅ Готово! Знайдено і форматовано як 2D-math: {math_count} формул!")

except Exception as e:
    print(f"Сталася помилка у Word: {e}")

finally:
    try:
        doc_word.Close()
        word.Quit()
    except:
        pass
